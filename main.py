# main.py — Refactored async FastAPI application

import asyncio
import io
import json
import logging
import os
import re
import time
from datetime import datetime, timezone, timedelta
from pathlib import Path
from fastapi import FastAPI, Request, Form, Depends, HTTPException, UploadFile, File, Body
from fastapi.responses import HTMLResponse, RedirectResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from sqlalchemy import inspect, text
from sqlalchemy.orm import Session
from passlib.context import CryptContext
try:
    from sse_starlette.sse import EventSourceResponse
except Exception:
    EventSourceResponse = None
    # Defer import failure until SSE endpoints are used; log a warning at runtime
from database import Base, engine, SessionLocal
from models import (
    User,
    InterviewAttempt,
    SkillProgress,
    UserProfile,
    Interview,
    UserSkillProfileRow,
    Course,
    Module,
    ModuleAttempt,
)
from user_skill_profile import (
    BasicUserInfo,
    ResumeData,
    UserSkillProfile,
    UserSkillVector,
    InterviewRecord,
    ScoreBreakdown,
    create_user_profile,
    detect_weaknesses,
    recommend_micro_courses,
    update_skill_score,
    update_skill_vector,
    calculate_overall_score,
    record_interview_result,
)
from core.llm.llm_service import LLMService
from core.prompts.prompt_manager import PromptManager
from services.rag.retriever import get_retriever
from services.rag.rag_pipeline import get_or_create_rag_pipeline, initialize_rag
from core.chains.base_chain import InterviewQuestionChain, EvaluationChain, SummaryChain
from speech.transcription import (
    transcribe_audio,
    analyze_speech_delivery,
    compute_confidence_score,
    compute_overall_score,
    compute_recruiter_verdict,
)
from utils.rl_helpers import (
    get_state_id,
    calculate_reward,
)
from services.rl.rl_service import ContextualBandit, INTERVIEW_ACTIONS, COURSE_ACTIONS
from models import QTable, UserState

Base.metadata.create_all(bind=engine)


def ensure_database_schema(engine):
    inspector = inspect(engine)
    table_names = inspector.get_table_names()

    with engine.begin() as conn:
        if "courses" in table_names:
            course_cols = {col["name"] for col in inspector.get_columns("courses")}
            if "role" not in course_cols:
                conn.execute(text('ALTER TABLE courses ADD COLUMN role VARCHAR;'))
            if "created_at" not in course_cols:
                conn.execute(text('ALTER TABLE courses ADD COLUMN created_at TIMESTAMP;'))
            if "updated_at" not in course_cols:
                conn.execute(text('ALTER TABLE courses ADD COLUMN updated_at TIMESTAMP;'))

        if "modules" in table_names:
            module_cols = {col["name"] for col in inspector.get_columns("modules")}
            if "is_final" not in module_cols:
                conn.execute(text('ALTER TABLE modules ADD COLUMN is_final BOOLEAN DEFAULT FALSE;'))


ensure_database_schema(engine)

import shutil

UPLOAD_DIR = "uploads"

# New architecture instances
llm_service = LLMService()
prompt_manager = PromptManager()
retriever = get_retriever()

# RAG Pipeline - initialized on startup
rag_pipeline = None

question_chain = InterviewQuestionChain(llm_service, prompt_manager, retriever)
evaluation_chain = EvaluationChain(llm_service, prompt_manager, retriever)
summary_chain = SummaryChain(llm_service, prompt_manager, retriever)

# ------------------------------------------------------------------
# Helper Functions
# ------------------------------------------------------------------
def categorize_weak_topics(weak_topics: list[str]) -> dict[str, list[str]]:

    technical = []
    communication = []

    for topic in weak_topics:
        t = topic.lower()

        if any(k in t for k in [
            "sql","database","python","algorithm","data structure",
            "machine learning","statistics","api","system design"
        ]):
            technical.append(topic)

        elif any(k in t for k in [
            "explain","clarity","structure","communication",
            "example","confidence","detail","depth"
        ]):
            communication.append(topic)

        else:
            # fallback
            technical.append(topic)

    return {
        "Technical Skills": list(set(technical)),
        "Communication & Answer Quality": list(set(communication))
    }

# ===========================================================================
# APP INIT
# ===========================================================================
app = FastAPI()
BASE_DIR = Path(__file__).parent
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")

logger = logging.getLogger(__name__)

# Use PBKDF2-SHA512 as primary (no 72-byte limit), keep bcrypt for backward compatibility
pwd_context = CryptContext(
    schemes=["pbkdf2_sha512", "bcrypt"],
    deprecated="auto"
)

# In-memory interview session store (swap for Redis in production)
interview_sessions: dict[str, dict] = {}

# In-memory resume text store (swap for Redis/DB in production)
resume_store: dict[str, str] = {}

# In-memory report store (latest report per user)
report_store: dict[str, dict] = {}


# ===========================================================================
# RAG PIPELINE STARTUP EVENT
# ===========================================================================
@app.on_event("startup")
async def startup_rag_pipeline():
    """Initialize RAG pipeline on application startup."""

    global rag_pipeline
    logger.info("Starting RAG pipeline initialization...")

    try:
        rag_pipeline = get_or_create_rag_pipeline()
        asyncio.create_task(initialize_rag_async())
        logger.info("RAG pipeline core created; async initialization task started")
    except Exception:
        logger.exception("RAG pipeline startup failed")
        rag_pipeline = None


async def initialize_rag_async():
    """Initialize RAG asynchronously."""
    if rag_pipeline is None:
        logger.warning("RAG async initialization skipped because the pipeline instance is unavailable")
        return

    try:
        await initialize_rag()
        logger.info("RAG pipeline initialized successfully")
    except Exception:
        logger.exception("RAG pipeline async initialization failed")


# ===========================================================================
# AI SERVICE HELPERS
# ===========================================================================

def extract_json(text: str):
    import json
    import re

    if not text:
        raise ValueError("Empty LLM response")

    # 🔹 Remove markdown code blocks
    text = text.strip()
    text = re.sub(r"^```json", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^```", "", text)
    text = re.sub(r"```$", "", text)

    # 🔹 Extract JSON object
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if not match:
        raise ValueError("No JSON object found in response")

    json_str = match.group(0)

    # 🔹 Remove invalid control characters (keep only valid JSON whitespace)
    json_str = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', json_str)

    # 🔹 Remove trailing commas (LLM bug)
    json_str = re.sub(r",\s*}", "}", json_str)
    json_str = re.sub(r",\s*]", "]", json_str)

    return json.loads(json_str, strict=False)


def safe_json_loads(json_str: str):
    """Safely parse JSON string by removing invalid control characters."""
    import json
    import re
    if not json_str:
        return None
    # Remove invalid control characters
    cleaned = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', json_str)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        return None


def normalize_questions_output(raw: str, count: int = 5) -> list[str]:
    questions = []

    if raw:
        try:
            parsed = extract_json(raw)
            if isinstance(parsed, list):
                questions = parsed
            elif isinstance(parsed, dict):
                if "questions" in parsed:
                    questions = parsed["questions"]
                elif "question" in parsed:
                    questions = parsed["question"]
                else:
                    for value in parsed.values():
                        if isinstance(value, list):
                            questions = value
                            break
        except Exception:
            # Fall back to regex and heuristic extraction
            matches = re.findall(r'"question"\s*:\s*"([^"]+)"', raw)
            if matches:
                questions = matches
            else:
                questions = [
                    line.strip().lstrip("0123456789.)- ")
                    for line in raw.splitlines()
                    if line.strip().endswith("?")
                ]

    cleaned_questions = []
    for q in questions:
        if isinstance(q, str):
            cleaned_questions.append(q.strip())
        elif isinstance(q, dict):
            if "question" in q:
                cleaned_questions.append(str(q["question"]).strip())
            elif "text" in q:
                cleaned_questions.append(str(q["text"]).strip())

    return [q for q in cleaned_questions if q][:count]


async def generate_interview_questions(role: str, level: str, count: int = 5, resume_text: str = "") -> list[str]:
    result = await question_chain.invoke(
        {
            "role": role,
            "level": level,
            "count": count,
            "resume_text": resume_text,
        }
    )

    if result.status != "success":
        return []

    try:
        return normalize_questions_output(result.output, count)
    except Exception:
        return []


async def evaluate_content(role: str, level: str, questions_answers: list) -> dict:
    answers = []
    weak_topics = []

    for qa in questions_answers:
        question = qa.get("question", "")
        answer = qa.get("answer", "")

        result = await evaluation_chain.invoke(
            {
                "role": role,
                "level": level,
                "question": question,
                "answer": answer,
            }
        )

        parsed = {}
        try:
            parsed = safe_json_loads(result.output) or {}
            if not isinstance(parsed, dict):
                parsed = {}
        except Exception:
            parsed = {}

        # Validate and extract score
        score = parsed.get("score", 0)
        try:
            score = float(score) if score else 0
        except (TypeError, ValueError):
            score = 0
        score = max(0, min(100, score))

        # Validate and extract lists
        answer_normalized = str(answer or "").strip().lower()
        if answer_normalized in ["", "(skipped)", "(no response)"]:
            strengths = ["No answer provided."]
        else:
            strengths = parsed.get("strengths", ["Answer attempted."])
            if not isinstance(strengths, list):
                strengths = ["Answer attempted."]
            strengths = [str(s).strip() for s in strengths if s][:3]
            if not strengths:
                strengths = ["Answer attempted."]

        weaknesses = parsed.get("weaknesses", ["Needs improvement."])
        if not isinstance(weaknesses, list):
            weaknesses = ["Needs improvement."]
        weaknesses = [str(w).strip() for w in weaknesses if w][:3]
        if not weaknesses:
            weaknesses = ["Needs improvement."]

        # Validate ideal_answer
        ideal_answer = parsed.get("ideal_answer", "Ideal answer unavailable.")
        if not isinstance(ideal_answer, str):
            ideal_answer = str(ideal_answer) if ideal_answer else "Ideal answer unavailable."

        # Validate weak_topics - MUST be list
        weak_topics_raw = parsed.get("weak_topics", [])
        if isinstance(weak_topics_raw, str):
            weak_topics_raw = [weak_topics_raw]
        if not isinstance(weak_topics_raw, list):
            weak_topics_raw = []
        weak_topics_list = [str(t).strip().lower() for t in weak_topics_raw if t]
        weak_topics_list = list(set(weak_topics_list))[:5]
        
        # ============================================================
        # RL: EXTRACT CKFS METRICS (NEW)
        # ============================================================
        rl_metrics = {
            "C": max(0.0, min(1.0, float(parsed.get("C", 0.0)))),
            "K": max(0.0, min(1.0, float(parsed.get("K", 0.0)))),
            "F": max(0.0, min(1.0, float(parsed.get("F", 0.0)))),
            "S": max(0.0, min(1.0, float(parsed.get("S", 0.0))))
        }
        
        item = {
            "score": score,
            "strengths": strengths,
            "weaknesses": weaknesses,
            "ideal_answer": ideal_answer,
            "weak_topics": weak_topics_list,
            "_rl_metrics": rl_metrics,  # Internal use only, not sent to frontend
        }
        
        weak_topics.extend(weak_topics_list)
        answers.append(item)

    # Deduplicate and clean weak_topics
    weak_topics_final = list(set([t.strip().lower() for t in weak_topics if t]))
    weak_topics_final = [t for t in weak_topics_final if t][:10]

    scores = [a["score"] for a in answers] if answers else [0]
    avg_score = sum(scores) / len(scores) if scores else 0

    return {
        "answers": answers,
        "weak_topics": weak_topics_final,
        "overall_feedback": "",
        "aggregate": {
            "technical_score": round(avg_score),
            "communication_score": round(avg_score * 0.9),
            "overall_score": round(avg_score),
        },
    }


async def evaluate_answers(role: str, questions_answers: list) -> dict:
    per_question = []

    for qa in questions_answers:
        question = qa.get("question", "")
        answer = qa.get("answer", "")

        result = await evaluation_chain.invoke(
            {
                "role": role,
                "level": "Junior",
                "question": question,
                "answer": answer,
            }
        )

        parsed = {}
        try:
            parsed = safe_json_loads(result.output) or {}
        except Exception:
            pass

        per_question.append({
            "question": question,
            "candidate_answer": answer,
            "feedback": parsed.get("strengths", []),
            "improved_answer": parsed.get("ideal_answer", ""),
        })

    return {
        "feedback_per_question": per_question,
        "improvement_tips": [
            "Practice structuring your answers with clear examples.",
            "Use the STAR method for behavioral responses.",
            "Focus on clarity, confidence, and relevance to the role.",
        ],
        "learning_resources": [
            {"topic": "Interview Preparation", "resource": "Practice common interview questions and structure answers clearly."}
        ],
    }


async def generate_performance_summary(report_data: dict) -> str:
    result = await summary_chain.invoke(
        {
            "role": report_data.get("candidate_profile", {}).get("role", "Candidate"),
            "score": report_data.get("overall_score", 0),
            "weak_topics": report_data.get("weak_topics", []),
            "attempted": len([
                a for a in report_data.get("detailed_answers", [])
                if a.get("transcript") not in ["", "(skipped)", "(no response)"]
            ]),
        }
    )

    return result.output or "Summary unavailable."


async def generate_resume_skill_profile(resume_text: str, role: str, level: str) -> dict:
    prompt = prompt_manager.get_prompt(
        "resume_skill_profile",
        resume_text=resume_text,
        role=role,
        experience_level=level,
    )

    raw = await llm_service.invoke(prompt, use_cache=False, json_mode=True)
    try:
        return extract_json(raw)
    except Exception:
        return {
            "identified_skills": [],
            "skill_gaps": [],
            "missing_for_role": [],
            "strength_percentage": 0,
            "improvement_suggestions": "Unable to parse resume insights.",
        }


# ===========================================================================
# DATABASE DEPENDENCY
# ===========================================================================
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# ===========================================================================
# SPEECH ANALYSIS HELPERS
# ===========================================================================
def analyze_speech_metrics(answer: str, duration_seconds: int) -> dict:
    """Analyze speech delivery metrics from answer text and duration."""
    words = answer.split()
    word_count = len(words)
    duration_minutes = duration_seconds / 60.0
    
    # Speaking pace (words per minute)
    wpm = word_count / duration_minutes if duration_minutes > 0 else 0
    
    # Filler word detection
    filler_words = ["um", "uh", "like", "you know", "so", "well", "actually", "basically", "literally", "totally"]
    filler_count = sum(1 for word in words if word.lower().strip('.,!?') in filler_words)
    filler_rate = filler_count / word_count if word_count > 0 else 0
    
    # Sentence clarity (simple metrics)
    sentences = re.split(r'[.!?]+', answer)
    sentences = [s.strip() for s in sentences if s.strip()]
    sentence_count = len(sentences)
    avg_words_per_sentence = word_count / sentence_count if sentence_count > 0 else 0
    
    # Clarity score (arbitrary, higher is better)
    clarity_score = min(10, max(0, 10 - (avg_words_per_sentence - 15) * 0.5 - filler_rate * 20))
    
    return {
        "word_count": word_count,
        "duration_seconds": duration_seconds,
        "speaking_pace_wpm": round(wpm, 1),
        "filler_word_count": filler_count,
        "filler_rate": round(filler_rate, 3),
        "sentence_count": sentence_count,
        "avg_words_per_sentence": round(avg_words_per_sentence, 1),
        "sentence_clarity_score": round(clarity_score, 1)
    }

# ===========================================================================
# AUTH HELPERS
# ===========================================================================
def get_current_user(request: Request, db: Session):
    username = request.cookies.get("user")
    if not username:
        return None
    return db.query(User).filter(User.username == username).first()

def hash_password(password: str):
    # PBKDF2-SHA512 has no length limit; full password is always hashed
    return pwd_context.hash(password)

def verify_password(plain: str, hashed: str):
    try:
        # Verify against hashed password (works with both PBKDF2 and legacy bcrypt)
        return pwd_context.verify(plain, hashed)
    except Exception as e:
        logger.warning(f"Password verification error: {e}")
        return False


def get_or_create_user_profile(db: Session, user: User) -> UserProfile:
    """
    Ensure a UserProfile row exists for the given auth user.

    This keeps creation logic in one place and can be reused from
    login, resume upload, and interview flows.
    """
    profile = db.query(UserProfile).filter(UserProfile.user_id == user.id).first()
    if profile:
        return profile

    profile = UserProfile(
        user_id=user.id,
        email=user.username,  # treat username as email by default
    )
    db.add(profile)
    db.commit()
    db.refresh(profile)
    return profile


# ---------------------------------------------------------------------------
# User Skill Profile helpers (DB-backed skill vector)
# ---------------------------------------------------------------------------

def get_skill_profile(db: Session, user_id: int):
    """Fetch the user skill profile row, or None if it doesn't exist."""
    return db.query(UserSkillProfileRow).filter(UserSkillProfileRow.user_id == user_id).first()


def create_skill_profile(db: Session, user_id: int) -> UserSkillProfileRow:
    """Create a new skill profile row for a user (no-op if already exists)."""
    profile = get_skill_profile(db, user_id)
    if profile:
        return profile

    profile = UserSkillProfileRow(
        user_id=user_id,
        technical_skills={},
        interview_skills={},
        communication_skills={},
        overall_score=0.0,
        interview_count=0,
    )
    db.add(profile)
    db.commit()
    db.refresh(profile)
    return profile
def update_skill_profile(db: Session, user_id: int, skill_data: dict):
    profile = get_skill_profile(db, user_id)

    if not profile:
        profile = create_skill_profile(db, user_id)

    if "technical_skills" in skill_data:
        profile.technical_skills = skill_data["technical_skills"]

    if "interview_skills" in skill_data:
        profile.interview_skills = skill_data["interview_skills"]

    if "communication_skills" in skill_data:
        profile.communication_skills = skill_data["communication_skills"]

    if "overall_score" in skill_data:
        profile.overall_score = float(skill_data["overall_score"])

    if "interview_count" in skill_data:
        profile.interview_count = int(skill_data["interview_count"])

    profile.last_updated = datetime.utcnow()

    db.add(profile)
    db.commit()
    db.refresh(profile)

    return profile
# ===========================================================================
# PAGE ROUTES
# ===========================================================================
@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    return templates.TemplateResponse(request, "home.html", {"request": request})

@app.get("/signup", response_class=HTMLResponse)
def signup_page(request: Request):
    return templates.TemplateResponse(request, "signup.html", {"request": request})

@app.post("/signup")
def signup(request: Request, username: str = Form(...), password: str = Form(...), db: Session = Depends(get_db)):
    if db.query(User).filter(User.username == username).first():
        return templates.TemplateResponse(request, "signup.html", {"request": request, "message": "User already exists"})
    hashed_password = hash_password(password)
    user = User(username=username, password=hashed_password)
    db.add(user)
    db.commit()
    return RedirectResponse("/login", status_code=303)

@app.get("/login", response_class=HTMLResponse)
def login_page(request: Request):
    return templates.TemplateResponse(request, "login.html", {"request": request})

@app.post("/login")
def login(request: Request, username: str = Form(...), password: str = Form(...), db: Session = Depends(get_db)):
    user = db.query(User).filter(User.username == username).first()

    if not user:
        return templates.TemplateResponse(request, "login.html", {"request": request, "message": "Invalid credentials"}
        )

    if not verify_password(password, user.password):
        return templates.TemplateResponse(request, "login.html", {"request": request, "message": "Invalid credentials"}
        )

    # Ensure a profile row exists for this user.
    get_or_create_user_profile(db, user)

    resp = RedirectResponse("/index", status_code=303)
    resp.set_cookie(key="user", value=username, httponly=True)
    return resp

@app.get("/logout")
def logout(request: Request, db: Session = Depends(get_db)):
    resp = RedirectResponse("/")
    resp.delete_cookie("user")
    return resp

@app.post("/update-resume")
def update_resume(request: Request, file: UploadFile, db: Session = Depends(get_db)):

    user = get_current_user(request, db)

    if not user:
        return RedirectResponse("/login", status_code=303)

    os.makedirs(UPLOAD_DIR, exist_ok=True)

    file_location = os.path.join(UPLOAD_DIR, f"{user.id}_{file.filename}")

    with open(file_location, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    profile = db.query(UserProfile).filter(UserProfile.user_id == user.id).first()

    if profile:
        profile.resume_file_path = file_location
        db.commit()

    return RedirectResponse("/profile", status_code=303)

@app.get("/index", response_class=HTMLResponse)
def index(request: Request, db: Session = Depends(get_db)):

    user = get_current_user(request, db)
    if not user:
        return RedirectResponse("/login")

    profile = db.query(UserProfile).filter(UserProfile.user_id == user.id).first()

    role = profile.role_applied_for if profile else ""
    level = profile.current_designation if profile else ""

    return templates.TemplateResponse(request, "index.html", {
            "request": request,
            "username": user.username,
            "saved_role": role,
            "saved_level": level,
        },
    )

@app.get("/progress", response_class=HTMLResponse)
def progress_page(request: Request, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user:
        return RedirectResponse("/login")
    skills = db.query(SkillProgress).filter(SkillProgress.user_id == user.id).all()
    return templates.TemplateResponse(request, "progress.html", {"request": request, "username": user.username, "skills": skills})


@app.get("/progress/", response_class=HTMLResponse)
def progress_page_slash(request: Request, db: Session = Depends(get_db)):
    """Support both /progress and /progress/ URLs."""
    return progress_page(request, db)

# ===========================================================================
# RESUME UPLOAD
# ===========================================================================
@app.post("/api/upload_resume")
async def upload_resume(
    request: Request,
    file: UploadFile = File(...),
    role: str = Form(...),
    level: str = Form(...),
    db: Session = Depends(get_db),
):
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")

    content = await file.read()
    original_name = file.filename or "resume"
    filename = original_name.lower()
    text = ""

    if filename.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    elif filename.endswith(".docx"):
        from docx import Document
        doc = Document(io.BytesIO(content))
        text = "\n".join(p.text for p in doc.paragraphs)
    else:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")

    text = text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="Could not extract text from the file")

    # Persist raw text in memory for question generation
    resume_store[user.username] = text

    # Persist file to disk for later download from profile
    upload_dir = Path("uploads")
    upload_dir.mkdir(parents=True, exist_ok=True)
    safe_name = f"user_{user.id}_{int(time.time())}_{original_name}"
    file_path = upload_dir / safe_name
    with open(file_path, "wb") as f:
        f.write(content)

    # Create / update structured skill profile using the resume
    profile_row = get_or_create_user_profile(db, user)

    # Ask LLM for compact skill profile using the new prompt manager
    try:
        parsed = await generate_resume_skill_profile(text, role, level)
    except Exception:
        parsed = {}

    skills = parsed.get("skills") if isinstance(parsed, dict) else None
    if not isinstance(skills, list):
        skills = []
    skills = [str(s).strip() for s in skills if str(s).strip()]

    skill_gaps = parsed.get("skill_gaps") if isinstance(parsed, dict) else None
    if not isinstance(skill_gaps, list):
        skill_gaps = []
    skill_gaps = [str(s).strip() for s in skill_gaps if str(s).strip()]

    improvement_suggestions = (
        parsed.get("improvement_suggestions") if isinstance(parsed, dict) else ""
    ) or ""

    strength_pct = 0.0
    if isinstance(parsed, dict):
        try:
            strength_pct = float(parsed.get("strength_percentage", 0.0))
        except (TypeError, ValueError):
            strength_pct = 0.0

    # Build Pydantic UserSkillProfile
    basic_info = BasicUserInfo(
        user_id=str(user.id),
        name=user.username,
        target_role=role,
        experience_level=level,
        resume_file_path=str(file_path),
    )
    resume_data = ResumeData(skills=skills)
    profile_obj: UserSkillProfile = create_user_profile(basic_info, resume_data)

    # Detect weaknesses based on initial profile (may be empty at this stage)
    detect_weaknesses(profile_obj)

    profile_row.role_applied_for = role
    profile_row.current_designation = level
    profile_row.resume_file_path = str(file_path)
    profile_row.extracted_skills = json.dumps(skills)
    profile_row.skill_strength_percentage = strength_pct
    profile_row.skill_gaps = json.dumps(skill_gaps)
    profile_row.improvement_suggestions = improvement_suggestions
    profile_row.profile_json = profile_obj.json()
    profile_row.updated_at = datetime.utcnow()

    db.add(profile_row)
    db.commit()
    return {
        "status": "ok",
        "length": len(text),
        "preview": text[:200],
        "skills": skills,
        "skill_gaps": skill_gaps,
        "strength_percentage": strength_pct,
    }

# ===========================================================================
# AUDIO TRANSCRIPTION (Whisper)
# ===========================================================================
@app.post("/api/transcribe")
async def transcribe_endpoint(file: UploadFile = File(...)):
    """
    Receive recorded audio from the frontend and convert it to text using Whisper.
    """

    audio_bytes = await file.read()

    temp_path = "temp_audio_" + str(datetime.now().timestamp()) + ".wav"

    # Save audio temporarily
    with open(temp_path, "wb") as f:
        f.write(audio_bytes)

    # Run Whisper transcription
    text = transcribe_audio(temp_path)

    return {"transcript": text}
# ===========================================================================
# INTERVIEW — START (returns page for voice interview)
# ===========================================================================
def _build_interview_context(request: Request, user, role: str, level: str, course_id: int | None, db: Session):
    completed_modules = []
    course_topics = []

    if course_id is not None:
        course = db.query(Course).filter(Course.id == course_id).first()
        if course and course.user_id == user.id:
            modules = (
                db.query(Module)
                .filter(Module.course_id == course_id)
                .order_by(Module.order_index.asc())
                .all()
            )
            completed_modules = [m.title for m in modules if m.is_completed]
            course_topics = [m.title for m in modules if m.title]

    interview_sessions[user.username] = {
        "role": role,
        "level": level,
        "course_id": course_id,
        "questions": [],
        "answers": [],
        "completed_modules": completed_modules,
        "course_topics": course_topics,
        "categories": []
    }

    return templates.TemplateResponse(request, "interview.html", {
            "request": request,
            "username": user.username,
            "user_id": user.id,
            "role": role,
            "level": level,
            "course_id": course_id,
        },
    )

@app.post("/start_interview/", response_class=HTMLResponse)
def start_interview_page(
    request: Request,
    role: str = Form(...),
    level: str = Form(...),
    course_id: int | None = Form(None),
    db: Session = Depends(get_db),
):
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")

    return _build_interview_context(request, user, role, level, course_id, db)

@app.get("/interview/start", response_class=HTMLResponse)
def start_interview_get(
    request: Request,
    role: str | None = None,
    level: str | None = None,
    course_id: int | None = None,
    db: Session = Depends(get_db),
):
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")
    effective_role = (role or "").strip()
    effective_level = (level or "").strip()
    completed_modules: list[str] = []
    course_topics: list[str] = []

    if course_id is not None:
        course = db.query(Course).filter(Course.id == course_id).first()
        if course and course.user_id == user.id:
            effective_role = (course.role or effective_role or "").strip()
            effective_level = (course.level or effective_level or "").strip()
            modules = (
                db.query(Module)
                .filter(Module.course_id == course_id)
                .order_by(Module.order_index.asc())
                .all()
            )
            completed_modules = [m.title for m in modules if m.is_completed and m.title]
            course_topics = [m.title for m in modules if m.title][:8]

    profile = db.query(UserProfile).filter(UserProfile.user_id == user.id).first()
    if not effective_role and profile and profile.role_applied_for:
        effective_role = profile.role_applied_for
    if not effective_level and profile and profile.current_designation:
        effective_level = profile.current_designation

    if not effective_role:
        effective_role = "Software Engineer"
    if not effective_level:
        effective_level = "Junior"

    resume_text = (resume_store.get(user.username) or "").strip()
    if not resume_text and profile and profile.resume_file_path:
        resume_path = Path(profile.resume_file_path)
        if resume_path.exists():
            try:
                suffix = resume_path.suffix.lower()
                if suffix == ".pdf":
                    from pypdf import PdfReader
                    reader = PdfReader(str(resume_path))
                    resume_text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
                elif suffix == ".docx":
                    from docx import Document
                    doc = Document(str(resume_path))
                    resume_text = "\n".join(p.text for p in doc.paragraphs).strip()
                else:
                    resume_text = resume_path.read_text(encoding="utf-8", errors="ignore").strip()
            except Exception:
                resume_text = ""
    if resume_text:
        resume_store[user.username] = resume_text

    interview_sessions[user.username] = {
        "role": effective_role,
        "level": effective_level,
        "course_id": course_id,
        "questions": [],
        "answers": [],
        "completed_modules": completed_modules,
        "course_topics": course_topics,
        "categories": []
    }

    logger.info("Interview started")

    return templates.TemplateResponse(request, "interview.html", {
            "request": request,
            "username": user.username,
            "user_id": user.id,
            "role": effective_role,
            "level": effective_level,
            "course_id": course_id,
        },
    )

# ===========================================================================
# COURSE GENERATION PAGE
# ===========================================================================
@app.post("/generate_course/", response_class=HTMLResponse)
def generate_course_page(
    request: Request,
    role: str = Form(...),
    level: str = Form(...),
    db: Session = Depends(get_db),
):
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")

    return templates.TemplateResponse(request, "course.html", {
            "request": request,
            "username": user.username,
            "role": role,
            "level": level,
        },
    )

# ===========================================================================
# INTERVIEW API — Retrieve next interview question using conversation history
# ===========================================================================
@app.post("/api/interview/next_question")
async def api_interview_next_question(request: Request, db: Session = Depends(get_db)):

    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")

    body = await request.json()

    user_id = body.get("user_id")
    history = body.get("history", [])

    # -----------------------------
    # VALIDATION
    # -----------------------------
    if user_id is None:
        raise HTTPException(status_code=400, detail="Missing user_id")

    if str(user_id) != str(user.id) and str(user_id) != str(user.username):
        raise HTTPException(status_code=403, detail="Invalid user_id")

    if not isinstance(history, list):
        raise HTTPException(status_code=400, detail="history must be a list")

    # -----------------------------
    # EXTRACT MEMORY FROM HISTORY
    # -----------------------------
    previous_questions = [
        h.get("question", "").strip()
        for h in history
        if isinstance(h, dict) and h.get("question")
    ]

    # -----------------------------
    # SESSION SETUP
    # -----------------------------
    session = interview_sessions.get(user.username)

    if session is None:
        interview_sessions[user.username] = {
            "role": body.get("role", "Software Engineer"),
            "level": body.get("level", "Junior"),
            "questions": [],
            "answers": [],
            "categories": []
        }
        session = interview_sessions[user.username]

    # Ensure categories always exist
    if "categories" not in session:
        session["categories"] = []

    used_categories = session.get("categories", [])

    # -----------------------------
    # CONTEXT
    # -----------------------------
    role = session.get("role", body.get("role", "Software Engineer"))
    level = session.get("level", body.get("level", "Junior"))
    resume_text = resume_store.get(user.username, "")
    completed_modules = session.get("completed_modules", [])
    course_topics = session.get("course_topics", [])

    # -----------------------------
    # LLM CALL (FIXED INPUT)
    # -----------------------------
    course_id = session.get("course_id", body.get("course_id"))
    llm_payload = {
        "role": role,
        "level": level,
        "resume_text": resume_text,
        "completed_modules": completed_modules,
        "course_topics": course_topics,
        "previous_questions": previous_questions,
        "used_categories": used_categories
    }
    if course_id is not None:
        llm_payload["context_priority"] = "course_focused"
        llm_payload["instruction_override"] = "Generate mostly questions from course_topics and completed_modules. Only 20-30% can be resume or general role-based."
        logger.info("Course-focused interview context applied for user_id=%s course_id=%s", user.id, course_id)

    result = await question_chain.invoke(llm_payload)

    if result.status != "success":
        raise HTTPException(status_code=500, detail="Failed to generate next interview question")

    # -----------------------------
    # PARSE RESPONSE SAFELY
    # -----------------------------
    try:
        parsed = extract_json(result.output)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse LLM response: {str(e)}")

    if not isinstance(parsed, dict) or "question" not in parsed:
        raise HTTPException(status_code=500, detail="Invalid question response format")

    # -----------------------------
    # UPDATE SESSION STATE
    # -----------------------------
    if "category" in parsed and parsed["category"]:
        session["categories"].append(parsed["category"])

    session["questions"].append(parsed.get("question", ""))

    # -----------------------------
    # RETURN RESPONSE
    # -----------------------------
    return JSONResponse(content=parsed)

# ===========================================================================
# INTERVIEW API — Batch evaluate all answers after interview ends
# ===========================================================================
@app.post("/api/interview/evaluate")
async def api_interview_evaluate(request: Request, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")


    body = await request.json()
    
    if "questions_answers" in body:
        # ---------------------------------------------------------------
        # Full analysis pipeline (Features 1-8)
        # ---------------------------------------------------------------
        questions_answers = body.get("questions_answers", [])
        if not questions_answers:
            raise HTTPException(status_code=400, detail="No questions_answers provided")


        session = interview_sessions.get(user.username, {})
        role = body.get("role", session.get("role", "Software Developer"))
        level = body.get("level", session.get("level", "mid"))
        n = len(questions_answers)

        # -- FEATURE 2: Speech Delivery Analysis -----------------------
        speech_analyses: list[dict] = []
        for qa in questions_answers:
            # Compute duration from timestamps when available
            duration = float(qa.get("duration", 60))
            start_time = qa.get("start_time", "")
            end_time = qa.get("end_time", "")
            if start_time and end_time:
                try:
                    st = datetime.fromisoformat(start_time.replace("Z", "+00:00"))
                    et = datetime.fromisoformat(end_time.replace("Z", "+00:00"))
                    duration = max((et - st).total_seconds(), 1.0)
                except Exception:
                    pass
            speech_analyses.append(
                analyze_speech_delivery(qa.get("answer", ""), duration)
            )

        # Aggregate voice metrics
        avg_pace = sum(a["speaking_pace_wpm"] for a in speech_analyses) / max(n, 1)
        total_fillers = sum(a["filler_word_count"] for a in speech_analyses)
        avg_clarity = sum(a["clarity_score"] for a in speech_analyses) / max(n, 1)
        avg_engagement = sum(a["engagement_score"] for a in speech_analyses) / max(n, 1)

        # -- FEATURE 4: Confidence Score -------------------------------
        confidence = compute_confidence_score(speech_analyses)

        # -- FEATURE 3: Content Analysis (LLM) ------------------------
        content_result = await evaluate_content(role, level, questions_answers)
        content_answers = content_result.get("answers", [])
        aggregate = content_result.get("aggregate", {})

        content_scores = [a.get("score", 50) for a in content_answers]
        content_avg = sum(content_scores) / max(len(content_scores), 1)

        # -- FEATURE 6: Overall Score & Verdict ------------------------
        overall = compute_overall_score(content_avg, avg_clarity, avg_engagement,questions_answers)
        verdict = compute_recruiter_verdict(overall, role)

        # -- FEATURE 4 (cont.): Session Metadata -----------------------
        total_duration = sum(a["duration_seconds"] for a in speech_analyses)

        # -- FEATURE 1: Candidate Profile ------------------------------
        candidate_profile = {
            "role": role,
            "level": level,
            "interview_date": datetime.now(timezone.utc).astimezone(timezone(timedelta(hours=5, minutes=30))).strftime("%Y-%m-%d %H:%M IST"),
            "total_questions": n,
        }
        # -- FEATURE 7: Detailed Answer Report -------------------------

        detailed_answers: list[dict] = []

        for i, qa in enumerate(questions_answers):

            ca = content_answers[i] if i < len(content_answers) else {}
            sa = speech_analyses[i] if i < len(speech_analyses) else {}

            # Ensure question text is from input
            question_text = qa.get("question", "")
            answer_text = qa.get("answer", "")


            # Normalize strengths
            strengths = ca.get("strengths", [])

            if isinstance(strengths, str):
                strengths = [strengths]

            if not strengths:
                strengths = ["Answer attempted."]

            # Normalize weaknesses
            weaknesses = ca.get("weaknesses", [])

            if isinstance(weaknesses, str):
                weaknesses = [weaknesses]

            if not weaknesses:
                weaknesses = ["Needs improvement."]

            detailed_answers.append({
                "question": question_text,
                "transcript": answer_text,
                "score": ca.get("score", 50),
                "strengths": strengths,
                "weaknesses": weaknesses,
                "ideal_answer": ca.get("ideal_answer", "No ideal answer generated."),
                "weak_topics": ca.get("weak_topics", []),
               "voice_metrics": {
    "speaking_pace_wpm": sa.get("speaking_pace_wpm", 0),
    "filler_word_count": sa.get("filler_word_count", 0),
    "clarity_score": sa.get("clarity_score", 0),
    "engagement_score": sa.get("engagement_score", 0),
    "word_count": sa.get("word_count", 0),
}
            })

        # -- FEATURE 8: Assemble Report --------------------------------
        report = {
            "candidate_profile": candidate_profile,
            "overall_score": overall,
            "verdict": verdict,
            "performance_summary": "",  # filled below by LLM
            "voice_analysis": {
                "speaking_pace_wpm": round(avg_pace, 1),
                "total_filler_words": total_fillers,
                "clarity_score": round(avg_clarity),
                "engagement_score": round(avg_engagement),
                "confidence_score": confidence,
                "total_duration_seconds": round(total_duration, 1),
                "per_answer": speech_analyses,
            },
            "content_analysis": {
                "average_score": round(content_avg),

                # values used by report.html
                "relevance_score": aggregate.get("technical_score", round(content_avg)),
                "depth_score": aggregate.get("communication_score", round(content_avg * 0.85)),
                "star_method_score": aggregate.get("overall_score", round(content_avg)),
             },
            "detailed_answers": detailed_answers,
        }

        # -- FEATURE 5: Performance Summary ----------------------------
        # Use overall_feedback from content evaluation to avoid a second LLM call.
        # Fall back to a separate LLM call only when the evaluator didn't provide one.
        overall_feedback = content_result.get("overall_feedback", "")
        if overall_feedback and overall_feedback.strip():
            report["performance_summary"] = overall_feedback.strip()
        else:
            summary = await generate_performance_summary(report)
            report["performance_summary"] = summary

        # Expose weak topics at the top level for report.html
        report["weak_topics"] = categorize_weak_topics(content_result.get("weak_topics", []))

        # ============================================================
        # RL: BANDIT-BASED COURSE RECOMMENDATION & REWARD FEEDBACK
        # ============================================================
        try:
            weak_topics = content_result.get("weak_topics", [])
            if isinstance(weak_topics, str):
                weak_topics = [weak_topics]
            weak_topics = [str(t).strip() for t in weak_topics if str(t).strip()]

            # STEP 1: Compute state
            state_id = get_state_id(overall, len(weak_topics))
            logger.info(f"[BANDIT] Computing state: score={overall:.1f}, weak_count={len(weak_topics)}, state_id={state_id}")

            # STEP 2: Get user state for previous score and session tracking
            user_state = db.query(UserState).filter(UserState.user_id == user.id).first()
            previous_score = float(user_state.avg_score) if user_state and user_state.avg_score is not None else overall
            prev_state_id = user_state.state_id if user_state else state_id
            logger.info(f"[BANDIT] Previous avg_score={previous_score:.1f}, prev_state_id={prev_state_id}")

            # STEP 3: Select action using bandit
            course_learner = ContextualBandit(db=db, action_space="course")
            action = course_learner.select_action(state_id, user_state)
            logger.info(f"[BANDIT] Selected action={action} for state_id={state_id}")

            # STEP 4: STRICT COURSE MAPPING - Action determines topics and difficulty
            def get_fundamentals(topics_list):
                """Get core/fundamental topics (first 2)."""
                return topics_list[:2] if len(topics_list) > 0 else ["core concepts"]

            def get_related_topics(topics_list):
                """Get weak topics + expanded related topics."""
                if len(topics_list) == 0:
                    return ["core concepts"]
                # Combine weak topics with logical expansions
                expanded = topics_list.copy()
                if len(expanded) < 3:
                    # Add generic related concepts if weak topics < 3
                    expanded.extend(["best practices", "advanced patterns"][:3-len(expanded)])
                return expanded

            def get_advanced_topics(topics_list):
                """Get advanced/deep topics for high performers."""
                if len(topics_list) == 0:
                    return ["advanced concepts"]
                # Focus on deeper aspects of weak topics
                advanced = [f"advanced {t}" if "advanced" not in t.lower() else t for t in topics_list[:2]]
                advanced.extend(["system design", "optimization"])
                return advanced[:4]

            # MAP ACTION → TOPICS + DIFFICULTY
            if action == "revision":
                course_topics = get_fundamentals(weak_topics)
                course_difficulty = "easy"
                logger.info(f"[BANDIT] ACTION=revision → topics={course_topics}, difficulty={course_difficulty}")

            elif action == "easy":
                course_topics = get_fundamentals(weak_topics)
                course_difficulty = "easy"
                logger.info(f"[BANDIT] ACTION=easy → topics={course_topics}, difficulty={course_difficulty}")

            elif action == "mixed":
                course_topics = get_related_topics(weak_topics)
                course_difficulty = "medium"
                logger.info(f"[BANDIT] ACTION=mixed → topics={course_topics}, difficulty={course_difficulty}")

            elif action == "advanced":
                course_topics = get_advanced_topics(weak_topics)
                course_difficulty = "hard"
                logger.info(f"[BANDIT] ACTION=advanced → topics={course_topics}, difficulty={course_difficulty}")

            else:
                # Fallback
                course_topics = weak_topics
                course_difficulty = "medium"
                logger.warning(f"[BANDIT] Unknown action={action}, using weak_topics")

            # STEP 5: ALWAYS run course generation with explicit topics and difficulty
            new_course_id = None
            try:
                new_course_id = await create_course_internal(
                    user,
                    role,
                    level,
                    weak_topics,
                    action,
                    db,
                    topics=course_topics,
                    difficulty=course_difficulty,
                )
                logger.info(f"[BANDIT] Course created: id={new_course_id}, action={action}, topics={course_topics}, difficulty={course_difficulty}")
            except Exception as course_error:
                logger.warning(
                    f"[BANDIT] Course generation failed for action={action}: {str(course_error)}"
                )
                new_course_id = None

            # Fallback: always try to create a course if first attempt failed
            if not new_course_id:
                try:
                    logger.info(f"[BANDIT] Attempting fallback course with revision action")
                    new_course_id = await create_course_internal(
                        user,
                        role,
                        level,
                        weak_topics,
                        "revision",
                        db,
                        topics=get_fundamentals(weak_topics),
                        difficulty="easy",
                    )
                    logger.info(f"[BANDIT] Fallback course created: id={new_course_id}")
                except Exception as fallback_error:
                    logger.warning(f"[BANDIT] Fallback course generation also failed: {str(fallback_error)}")
                    new_course_id = None

            # STEP 6: Fetch course details for response
            new_course = None
            if new_course_id:
                new_course_row = db.query(Course).filter(Course.id == new_course_id).first()
                if new_course_row:
                    new_course = {
                        "course_id": new_course_row.id,
                        "title": new_course_row.title,
                    }

            # STEP 7: Calculate SIMPLE reward = score improvement normalized to [-1, 1]
            score_improvement = overall - previous_score  # range: -100 to +100
            reward = score_improvement / 100.0  # normalize to [-1, +1]
            reward = max(min(reward, 1.0), -1.0)  # clamp to [-1, 1]
            logger.info(f"[BANDIT] Reward calculation: overall={overall:.1f} - previous={previous_score:.1f} = improvement={score_improvement:.1f} → normalized_reward={reward:.3f}")

            # STEP 8: Update bandit with reward (no next_state needed)
            course_learner.update_action_value(prev_state_id, action, reward)
            logger.info(f"[BANDIT] Q-value updated: state={prev_state_id}, action={action}, reward={reward:.3f}")

            # STEP 9: Update user state for next session
            if user_state:
                user_state.state_id = state_id
                user_state.last_score = overall
                user_state.avg_score = (
                    user_state.avg_score * user_state.session_count + overall
                ) / (user_state.session_count + 1)
                user_state.current_proficiency = user_state.avg_score / 100.0
                user_state.weak_topics = weak_topics
                user_state.session_count += 1
            else:
                user_state = UserState(
                    user_id=user.id,
                    state_id=state_id,
                    weak_topics=weak_topics,
                    avg_score=overall,
                    last_score=overall,
                    current_proficiency=overall / 100.0,
                    session_count=1,
                )
                db.add(user_state)

            db.flush()

            # STEP 10: Populate report with course and action
            report["new_course_id"] = new_course_id
            if new_course is not None:
                report["new_course"] = new_course
            report["recommended_action"] = action
            report["rl_metrics"] = {
                "state": state_id,
                "previous_state": prev_state_id,
                "action": action,
                "course_topics": course_topics,
                "course_difficulty": course_difficulty,
                "reward": reward,
                "score_improvement": score_improvement,
                "new_course_id": new_course_id,
            }

            logger.info(
                f"[BANDIT] COMPLETE: state={state_id}, action={action}, reward={reward:.3f}, "
                f"course_id={new_course_id}, topics={course_topics}"
            )
        except Exception as e:
            logger.warning(f"[BANDIT] RL course recommendation failed: {str(e)}", exc_info=True)
            report["new_course_id"] = None
            report["recommended_action"] = "revision"

        # Store report for the /report page (latest only)
        report_store[user.username] = report
        interview_row = Interview(
            user_id=user.id,
            role=role,
            date=datetime.utcnow(),
            score=overall,
            report_json=json.dumps(report),
        )
        db.add(interview_row)
        db.commit()
        db.refresh(interview_row)

        # Update rich UserSkillProfile based on this interview
        profile_row = get_or_create_user_profile(db, user)
        profile_obj: UserSkillProfile
        try:
            if profile_row.profile_json:
                profile_obj = UserSkillProfile.model_validate_json(profile_row.profile_json)
            else:
                # Minimal fallback profile if none existed yet
                basic_info = BasicUserInfo(
                    user_id=str(user.id),
                    name=user.username,
                    target_role=role,
                    experience_level=level,
                    resume_file_path=profile_row.resume_file_path or "",
                )
                resume_data = ResumeData(skills=[])
                profile_obj = create_user_profile(basic_info, resume_data)
        except Exception:
            basic_info = BasicUserInfo(
                user_id=str(user.id),
                name=user.username,
                target_role=role,
                experience_level=level,
                resume_file_path=profile_row.resume_file_path or "",
            )
            resume_data = ResumeData(skills=[])
            profile_obj = create_user_profile(basic_info, resume_data)

        # Record interview per question into profile
        for i, qa in enumerate(questions_answers):
            ca = content_answers[i] if i < len(content_answers) else {}
            topic = role  # Treat role as the main topic/skill for now

            sb = ScoreBreakdown(
                correctness=float(ca.get("score", 50)),
                conceptual_depth=float(ca.get("score", 50)),
                clarity=float(avg_clarity),
                feedback="",
            )
            rec = InterviewRecord(
                question=qa.get("question", ""),
                topic=topic,
                answer_transcript=qa.get("answer", ""),
                evaluation_score=float(ca.get("score", 50)),
                score_breakdown=sb,
            )
            record_interview_result(profile_obj, rec)

        # Sync weak topics from LLM result - handle both dict and list
        weak_from_llm = []
        weak_topics_raw = content_result.get("weak_topics", [])
        if isinstance(weak_topics_raw, dict):
            for topics in weak_topics_raw.values():
                if isinstance(topics, list):
                    weak_from_llm.extend(topics)
                else:
                    weak_from_llm.append(str(topics) if topics else "")
        elif isinstance(weak_topics_raw, list):
            weak_from_llm.extend(weak_topics_raw)
        
        # Safe access to weak_topics with fallback
        profile_weak_topics = getattr(profile_obj, 'weak_topics', [])
        if not isinstance(profile_weak_topics, list):
            profile_weak_topics = []
        
        for w in weak_from_llm:
            if not w:
                continue
            w_str = str(w).strip() if w else ""
            if not w_str:
                continue
            w_norm = w_str.lower()
            if w_norm not in {t.lower() for t in profile_weak_topics if t}:
                profile_weak_topics.append(w_str)
        
        # Update profile with clean weak_topics list
        profile_obj.weak_topics = profile_weak_topics

        # Recompute weaknesses & recommend micro-courses
        detect_weaknesses(profile_obj)
        recommend_micro_courses(profile_obj)

        # --------- Update skill vector based on interview metrics ---------
        def _int(v, default=0):
            try:
                return int(round(float(v)))
            except Exception:
                return default

        voice = report.get("voice_analysis", {})
        content = report.get("content_analysis", {})

        interview_metrics = {
            "relevance": _int(content.get("relevance_score", 0)),
            "explanation_depth": _int(content.get("depth_score", 0)),
            "problem_solving": _int(content.get("average_score", 0)),
            "structured_thinking": _int(content.get("average_score", 0)),
            "star_method": _int(content.get("star_method_score", 0)),

            "clarity": _int(voice.get("clarity_score", 0)),
            "confidence": _int(voice.get("confidence_score", 0)),
            "engagement": _int(voice.get("engagement_score", 0)),
            "speaking_pace": _int(voice.get("speaking_pace_wpm", 0)),
            "filler_control": 100 - min(100, _int(voice.get("total_filler_words", 0))),
}

        # Update the stored skill vector using the interview metrics
        skill_vector = UserSkillVector(
            technical_skills=profile_obj.technical_skills,
            interview_skills=profile_obj.interview_skills,
            communication_skills=profile_obj.communication_skills,
        )
        skill_vector = update_skill_vector(skill_vector, interview_metrics)

        profile_obj.technical_skills = skill_vector.technical_skills
        profile_obj.interview_skills = skill_vector.interview_skills
        profile_obj.communication_skills = skill_vector.communication_skills
        profile_obj.overall_score = calculate_overall_score(skill_vector)

        # Aggregate simple fields for quick display in profile page
        skills_list = []
        avg_score = overall

        if hasattr(profile_obj, "skill_graph") and profile_obj.skill_graph:
            try:
                skills_list = [node.skill_name for node in profile_obj.skill_graph.values() if hasattr(node, 'skill_name')]
                if profile_obj.skill_graph:
                    scores = [n.score for n in profile_obj.skill_graph.values() if hasattr(n, 'score')]
                    if scores:
                        avg_score = sum(scores) / len(scores)
            except Exception:
                pass

        profile_row.role_applied_for = role
        profile_row.current_designation = level
        profile_row.extracted_skills = json.dumps(skills_list)
        profile_row.skill_strength_percentage = float(profile_obj.overall_score)
        profile_row.skill_gaps = json.dumps(getattr(profile_obj, 'weak_topics', []))
        # Use performance summary as latest improvement suggestions snapshot
        profile_row.improvement_suggestions = report.get("performance_summary", "")
        profile_row.profile_json = profile_obj.json()
        profile_row.updated_at = datetime.utcnow()

        # Persist a compact skill vector record for quick access / querying.
        update_skill_profile(
            db,
            user.id,
            {
                "technical_skills": profile_obj.technical_skills.dict(),
                "interview_skills": profile_obj.interview_skills.dict(),
                "communication_skills": profile_obj.communication_skills.dict(),
                "overall_score": profile_obj.overall_score,
                "interview_count": profile_obj.interview_count,
            },
        )

        # Return the updated skill profile as part of the report payload.
        report["skill_profile"] = profile_obj.dict()

        db.add(profile_row)

        # Persist answers to DB
        for i, qa in enumerate(questions_answers):
            attempt = InterviewAttempt(
                user_id=user.id,
                role=role,
                topic="voice-interview",
                difficulty=level,
                answer=qa.get("answer", ""),
                feedback="",
            )
            db.add(attempt)

        # Update skill progress
        skill = (
            db.query(SkillProgress)
            .filter(SkillProgress.user_id == user.id, SkillProgress.skill == role)
            .first()
        )
        if not skill:
            skill = SkillProgress(
                user_id=user.id, skill=role, attempts=1, weak=overall < 50
            )
            db.add(skill)
        else:
            skill.attempts += 1
            skill.weak = overall < 50

        db.commit()

        return report
    
    else:
        # Old format
        role = body.get("role", "")
        answers = body.get("answers", [])  # [{question, answer}, ...]

        if not answers:
            raise HTTPException(status_code=400, detail="No answers provided")

        # Store answers in session
        session = interview_sessions.get(user.username, {})
        session["answers"] = answers

        evaluation = await evaluate_answers(role, answers)

        # Normalize top-level keys (AI may use variant names)
        if isinstance(evaluation, dict):
            for alt in ("evaluations", "feedback", "responses", "results"):
                if alt in evaluation and "feedback_per_question" not in evaluation:
                    evaluation["feedback_per_question"] = evaluation.pop(alt)
                    break
            for alt in ("tips", "suggestions", "general_tips"):
                if alt in evaluation and "improvement_tips" not in evaluation:
                    evaluation["improvement_tips"] = evaluation.pop(alt)
                    break
            for alt in ("resources", "recommended_resources", "study_resources"):
                if alt in evaluation and "learning_resources" not in evaluation:
                    evaluation["learning_resources"] = evaluation.pop(alt)
                    break

        # Validate and build fallback if parsing failed or structure is wrong
        if not isinstance(evaluation, dict) or "feedback_per_question" not in evaluation:
            evaluation = {
                "feedback_per_question": [
                    {
                        "question": qa.get("question", ""),
                        "candidate_answer": qa.get("answer", ""),
                        "feedback": "",
                        "improved_answer": "",
                    }
                    for qa in answers
                ],
                "improvement_tips": [
                    "Practice structuring your answers with concrete examples.",
                    "Use the STAR method (Situation, Task, Action, Result) for behavioral questions.",
                    "Research the company and role thoroughly before interviews.",
                ],
                "learning_resources": [{"topic": "Interview Preparation", "resource": "Practice common behavioral and technical questions for your role."}],
            }

        # Normalize per-question field names (AI may use variants)
        for i, fb in enumerate(evaluation.get("feedback_per_question", [])):
            qa = answers[i] if i < len(answers) else {}
            for src, dst in [
                ("answer", "candidate_answer"),
                ("user_answer", "candidate_answer"),
                ("better_answer", "improved_answer"),
                ("suggested_answer", "improved_answer"),
                ("sample_answer", "improved_answer"),
                ("ideal_answer", "improved_answer"),
            ]:
                if src in fb and dst not in fb:
                    fb[dst] = fb.pop(src)
            fb.setdefault("question", qa.get("question", ""))
            fb.setdefault("candidate_answer", qa.get("answer", ""))
            fb.setdefault("feedback", "")
            fb.setdefault("improved_answer", "")

        # Normalize learning_resources to always be list of {topic, resource}
        raw_resources = evaluation.get("learning_resources", [])
        normalized = []
        for r in raw_resources:
            if isinstance(r, dict):
                topic = r.get("topic") or r.get("name") or r.get("title") or ""
                resource = r.get("resource") or r.get("description") or r.get("link") or r.get("url") or ""
                normalized.append({"topic": topic, "resource": resource})
            elif isinstance(r, str):
                normalized.append({"topic": r, "resource": ""})
        evaluation["learning_resources"] = normalized

        # Persist each answer to DB
        for i, qa in enumerate(answers):
            attempt = InterviewAttempt(
                user_id=user.id,
                role=role,
                topic="voice-interview",
                difficulty="adaptive",
                answer=qa.get("answer", ""),
                feedback="",
            )
            db.add(attempt)

        # Update skill progress
        skill = db.query(SkillProgress).filter(SkillProgress.user_id == user.id, SkillProgress.skill == role).first()
        if not skill:
            skill = SkillProgress(user_id=user.id, skill=role, attempts=1, weak=False)
            db.add(skill)
        else:
            skill.attempts += 1

        db.commit()

        return evaluation

# ===========================================================================
# REPORT PAGE
# ===========================================================================
@app.get("/report", response_class=HTMLResponse)
def report_page(request: Request, db: Session = Depends(get_db)):
    """Render the full interview performance report."""
    user = get_current_user(request, db)
    if not user:
        return RedirectResponse("/login")
    report = report_store.get(user.username)
    if not report:
        return RedirectResponse("/index")
    return templates.TemplateResponse(request, "report.html", {
            "request": request,
            "username": user.username,
            "report": report,
        },
    )


@app.get("/interview-report/{interview_id}", response_class=HTMLResponse)
def interview_report_page(
    request: Request,
    interview_id: int,
    db: Session = Depends(get_db),
):
    """
    View a past interview report stored in DB (secure).
    """

    from models import Interview
    import json

    user = get_current_user(request, db)
    if not user:
        return RedirectResponse("/login")

    # ✅ IMPORTANT: restrict to user's own reports
    interview = (
        db.query(Interview)
        .filter(
            Interview.id == interview_id,
            Interview.user_id == user.id
        )
        .first()
    )

    if not interview:
        raise HTTPException(status_code=404, detail="Report not found")

    try:
        report = json.loads(interview.report_json)
    except Exception:
        report = {}

    return templates.TemplateResponse(request, "report.html", {
            "request": request,
            "username": user.username,
            "report": report
        },
    )


@app.get("/profile", response_class=HTMLResponse)
def profile_page(request: Request, db: Session = Depends(get_db)):

    user = get_current_user(request, db)
    if not user:
        return RedirectResponse("/login")

    from models import Interview
    import json

    # Ensure profile exists
    profile_row = get_or_create_user_profile(db, user)

    # -------------------------------
    # DEFAULT SKILL PROFILE (fallback)
    # -------------------------------
    profile_data = {
        "technical_skills": {
            "dsa": 50,
            "dbms": 50,
            "operating_systems": 50,
            "computer_networks": 50,
            "system_design": 50,
        },
        "interview_skills": {
            "relevance": 50,
            "explanation_depth": 50,
            "structured_thinking": 50,
            "problem_solving": 50,
            "star_method": 50,
        },
        "communication_skills": {
            "clarity": 50,
            "confidence": 50,
            "engagement": 50,
            "speaking_pace": 50,
            "filler_control": 50,
        },
        "overall_score": 50
    }

    # -------------------------------
    # LOAD STORED SKILL PROFILE
    # -------------------------------
    if profile_row.profile_json:
        try:
            loaded = json.loads(profile_row.profile_json)
            profile_data.update(loaded)
        except Exception:
            pass

    # -------------------------------
    # EXTRACTED SKILLS
    # -------------------------------
    extracted_skills = []
    if profile_row.extracted_skills:
        try:
            extracted_skills = json.loads(profile_row.extracted_skills)
        except:
            extracted_skills = []

    # -------------------------------
    # SKILL GAPS
    # -------------------------------
    skill_gaps = []
    if profile_row.skill_gaps:
        try:
            skill_gaps = json.loads(profile_row.skill_gaps)
        except:
            skill_gaps = []

    # -------------------------------
    # FETCH INTERVIEW HISTORY
    # -------------------------------
    interviews = (
        db.query(Interview)
        .filter(Interview.user_id == user.id)
        .order_by(Interview.date.asc())
        .all()
    )

    # -------------------------------
    # GROUP BY ROLE (History section)
    # -------------------------------
    history_by_role = {}

    for iv in interviews:
       normalized_role = iv.role.strip().lower()   # ✅ normalize

       history_by_role.setdefault(normalized_role, []).append(iv)

    # -------------------------------
    # BUILD TIMELINE (for graphs)
    # -------------------------------
    timeline_by_role = {}

    for iv in interviews:
        role = iv.role.strip().lower()

        if role not in timeline_by_role:
            timeline_by_role[role] = []

        timeline_by_role[role].append({
            "date": iv.date.strftime("%Y-%m-%d"),
            "score": iv.score
        })

    # -------------------------------
    # IMPROVEMENT MESSAGE
    # -------------------------------
    improvement_message = (
        profile_row.improvement_suggestions
        if profile_row and profile_row.improvement_suggestions
        else "Keep practicing interviews to improve your profile."
    )

    # -------------------------------
    # COURSE HISTORY
    # -------------------------------
    course_history = []
    courses = (
        db.query(Course)
        .filter(Course.user_id == user.id)
        .order_by(Course.created_at.asc())
        .all()
    )
    for course in courses:
        modules = (
            db.query(Module)
            .filter(Module.course_id == course.id)
            .order_by(Module.order_index.asc())
            .all()
        )
        module_entries = []
        for module in modules:
            attempt = (
                db.query(ModuleAttempt)
                .filter(
                    ModuleAttempt.user_id == user.id,
                    ModuleAttempt.module_id == module.id
                )
                .order_by(ModuleAttempt.created_at.desc())
                .first()
            )
            score = attempt.score if attempt else None
            attempts = 0
            answers = []
            is_passed = False
            total_questions = attempt.total_questions if attempt else None
            if attempt:
                payload = attempt.answers
                if isinstance(payload, dict):
                    try:
                        attempts = int(payload.get("attempt_count", 1))
                    except Exception:
                        attempts = 1
                    is_passed = bool(payload.get("is_passed", False))
                    stored_answers = payload.get("answers", [])
                    answers = stored_answers if isinstance(stored_answers, list) else []
                elif isinstance(payload, list):
                    attempts = 1
                    answers = payload
                    is_passed = bool(score is not None and score >= 2)

            module_entries.append(
                {
                    "module_id": module.id,
                    "title": module.title,
                    "order_index": module.order_index,
                    "is_unlocked": module.is_unlocked,
                    "is_completed": module.is_completed,
                    "score": score,
                    "total_questions": total_questions,
                    "is_passed": is_passed,
                    "attempts": attempts,
                    "answers": answers,
                }
            )

        course_history.append(
            {
                "course_id": course.id,
                "title": course.title,
                "role": course.role,
                "level": course.level,
                "status": course.status,
                "created_at": course.created_at.strftime("%Y-%m-%d") if course.created_at else None,
                "modules": module_entries,
            }
        )

    logger.info("Course history built for user_id=%s courses=%s", user.id, len(course_history))

    return templates.TemplateResponse(request, "profile.html", {
            "request": request,
            "username": user.username,
            "user_profile": profile_row,
            "skill_profile": profile_data,
            "extracted_skills": extracted_skills,
            "skill_gaps": skill_gaps,
            "interview_history_by_role": history_by_role,
            "timeline_by_role": timeline_by_role,
            "improvement_message": improvement_message,
            "course_history": course_history
        },
    )

@app.get("/interview-history", response_class=HTMLResponse)
def interview_history_redirect(request: Request):
    """
    Simple semantic route that redirects to the profile page where
    interview history is rendered.
    """
    return RedirectResponse("/profile")


# Deprecated: replaced by DB-based course system
# @app.get("/api/course/stream")
# async def api_course_stream(request: Request, role: str, level: str, db: Session = Depends(get_db)):
#     user = get_current_user(request, db)
#     if not user:
#         raise HTTPException(status_code=401, detail="Not logged in")
#
#     course = Course(
#         user_id=user.id,
#         role=role,
#         title=f"{level.title()} {role} Course",
#         description="",
#         level=level,
#         status="draft"
#     )
#     db.add(course)
#     db.flush()
#
#     async def event_generator():
#         import asyncio
#
#         # Heartbeat
#         yield {"event": "heartbeat", "data": "connected"}
#
#         # =========================
#         # STAGE 1: OUTLINE
#         # =========================
#         yield {"event": "status", "data": "Generating course outline..."}
#
#         await asyncio.sleep(0.5)
#
#         outline_prompt_input = {
#             "skill": role,
#             "level": level,
#             "duration_hours": 20
#         }
#
#         # ===== OUTLINE =====
#         outline_prompt = prompt_manager.get_prompt("course_outline", **outline_prompt_input)
#         outline_text = await llm_service.invoke(outline_prompt)
#
#         try:
#             outline_json = extract_json(outline_text)
#         except:
#             match = re.search(r"\{.*\}", outline_text, re.DOTALL)
#             if match:
#                 outline_json = json.loads(match.group(0))
#             else:
#                 raise ValueError("Invalid outline JSON")
#
#         course.title = outline_json.get("course_title", course.title)
#         course.description = outline_json.get("description", "")
#         db.commit()
#
#         # Normalize modules and persist skeletons
#         modules = []
#         module_records = []
#         for idx, mod in enumerate(outline_json.get("modules", [])):
#             title = mod.get("module_title") or mod.get("title")
#             description = ", ".join(mod.get("topics", [])) if mod.get("topics") else mod.get("description", "")
#
#             module_record = Module(
#                 course_id=course.id,
#                 title=title,
#                 description=description,
#                 order_index=idx,
#                 is_unlocked=(idx == 0)
#             )
#             db.add(module_record)
#             db.flush()
#             module_records.append(module_record)
#
#             modules.append({
#                 "id": module_record.id,
#                 "title": title,
#                 "description": description
#             })
#
#         db.commit()
#
#         # Send outline
#         yield {
#             "event": "outline",
#             "data": json.dumps({
#                 "course_title": outline_json.get("course_title"),
#                 "description": outline_json.get("description"),
#                 "modules": modules
#             })
#         }
#
#         # =========================
#         # STAGE 2: MODULES (Deep Dive)
#         # =========================
#         for i, mod in enumerate(modules):
#             if await request.is_disconnected():
#                 return
#
#             yield {
#                 "event": "status",
#                 "data": f"Generating module {i+1}/{len(modules)}..."
#             }
#
#             module_prompt_input = {
#                 "skill": role,
#                 "module": mod["title"],
#                 "level": level
#             }
#
#             # ===== MODULE =====
#             module_prompt = prompt_manager.get_prompt("course_module_detail", **module_prompt_input)
#             module_text = await llm_service.invoke(module_prompt)
#
#             try:
#                 module_json = json.loads(module_text)
#             except:
#                 match = re.search(r"\{.*\}", module_text, re.DOTALL)
#                 if match:
#                     module_json = json.loads(match.group(0))
#                 else:
#                     raise ValueError("Invalid module JSON")
#
#             # Safety check
#             if len(module_json.get("quiz", [])) != 3:
#                 module_json["quiz"] = []
#
#             module_record = module_records[i]
#             module_record.content = module_json.get("content_markdown", "")
#             module_record.quiz = module_json.get("quiz", [])
#             db.commit()
#
#             yield {
#                 "event": "module_detail",
#                 "data": json.dumps({
#                     "module_id": module_record.id,
#                     "index": i,
#                     "module_title": module_json.get("module_title", module_record.title),
#                     "content_markdown": module_record.content,
#                     "quiz": module_record.quiz,
#                     "external_practice_tasks": module_json.get("external_practice_tasks", [])
#                 })
#             }
#
#         # =========================
#         # DONE
#         # =========================
#         yield {"event": "done", "data": "Course generation complete!"}
#
#     return EventSourceResponse(event_generator())
# ===========================================================================
# SKELETON-FIRST COURSE GENERATION SYSTEM (NEW ARCHITECTURE)
# ===========================================================================

@app.post("/api/course/create")
async def create_course(
    request: Request,
    role: str = Form(...),
    level: str = Form(...),
    db: Session = Depends(get_db)
):
    """
    Create course skeleton (outline only, no module content yet).
    
    - Generate outline using PromptManager
    - Create Course record
    - Create Module records (title + description only)
    - Unlock first module
    - Return course_id and modules list
    """
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")
    
    try:
        # Generate outline using LLM
        outline_prompt_input = {
            "skill": role,
            "level": level,
            "duration_hours": 20
        }
        
        # ===== OUTLINE =====
        outline_prompt = prompt_manager.get_prompt("course_outline", **outline_prompt_input)
        outline_text = await llm_service.invoke(outline_prompt)
        
        # Parse LLM response (extract JSON)
        try:
            outline_json = extract_json(outline_text)

        except Exception as e:

            match = re.search(r"\{.*\}", outline_text, re.DOTALL)

            if match:
                try:
                   outline_json = json.loads(match.group(0))
                except Exception:
                    raise HTTPException(status_code=500, detail="Invalid outline JSON format")
            else:
                raise HTTPException(status_code=500, detail="No JSON found in outline response")
        
        # Create Course record
        course = Course(
            user_id=user.id,
            role=role,
            title=outline_json.get("course_title", f"{level.title()} {role} Course"),
            description=outline_json.get("description", ""),
            level=level,
            status="generated"
        )
        db.add(course)
        db.flush()  # Get course.id
        
        # Create Module records (skeleton only, no content yet)
        modules_data = []
        outline_modules = outline_json.get("modules", [])
        total_modules = len(outline_modules)

        for idx, mod_data in enumerate(outline_modules):
            title = mod_data.get("module_title") or mod_data.get("title") or f"Module {idx+1}"
            description = ", ".join(mod_data.get("topics", [])) if mod_data.get("topics") else mod_data.get("description", "")
            is_final_module = idx == total_modules - 1

            module = Module(
                course_id=course.id,
                title=title,
                description=description,
                order_index=idx,
                is_unlocked=(idx == 0),
                is_final=is_final_module
            )

            db.add(module)
            db.flush()

            modules_data.append({
                "id": module.id,
                "title": module.title,
                "description": module.description,
                "order_index": module.order_index,
                "is_unlocked": module.is_unlocked,
                "is_completed": module.is_completed,
                "is_final": module.is_final
            })

        # Ensure minimum 5 modules
        if total_modules < 5:
            for i in range(total_modules, 5):
                title = f"Additional Module {i+1}"
                description = "Extra learning module"
                is_final_module = i == 4

                module = Module(
                    course_id=course.id,
                    title=title,
                    description=description,
                    order_index=i,
                    is_unlocked=(i == 0 and total_modules == 0),
                    is_final=is_final_module
                )

                db.add(module)
                db.flush()

                modules_data.append({
                    "id": module.id,
                    "title": module.title,
                    "description": module.description,
                    "order_index": module.order_index,
                    "is_unlocked": module.is_unlocked,
                    "is_completed": module.is_completed,
                    "is_final": module.is_final
                })

        # ✅ COMMIT AFTER LOOP
        db.commit()
        logger.info("Course created")

        # ✅ RETURN AFTER LOOP
        return RedirectResponse(url=f"/course/{course.id}", status_code=303)
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Course creation failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to create course")


async def create_course_internal(
    user,
    role: str,
    level: str,
    weak_topics: list[str],
    action: str,
    db: Session,
    topics: list[str] = None,
    difficulty: str = "medium",
) -> int | None:
    """
    Create a course with explicit topic selection and difficulty.
    
    Args:
        user: User object
        role: Job role
        level: Experience level
        weak_topics: Original weak topics from interview
        action: Bandit action (revision, easy, mixed, advanced)
        db: Database session
        topics: EXPLICIT topics to cover (if None, uses weak_topics)
        difficulty: EXPLICIT difficulty level (easy, medium, hard)
    """
    weak_topics = [str(t).strip() for t in weak_topics if str(t).strip()]
    weak_topics = weak_topics or ["core concepts"]
    
    # Use explicit topics if provided, otherwise use weak_topics
    course_topics = topics if topics else weak_topics
    course_topics = [str(t).strip() for t in course_topics if str(t).strip()]
    topics_text = ", ".join(course_topics)

    prompt = f"""
Generate a course outline.

COURSE PARAMETERS:
- Action: {action}
- Difficulty Level: {difficulty}
- Topics to Cover: {topics_text}

STRICT REQUIREMENTS:
- FIRST 2 modules MUST directly cover these topics: {topics_text}
- Each topic MUST appear explicitly in module titles
- Set difficulty to {difficulty.upper()} throughout the course
- Do NOT generalize or skip required topics
- Remaining modules can expand related concepts
- Module difficulty should match: {difficulty}
"""

    outline_prompt_input = {
        "skill": role,
        "level": level,
        "duration_hours": 20,
        "strict_requirements": prompt,
    }

    outline_prompt = prompt_manager.get_prompt("course_outline", **outline_prompt_input)
    outline_text = await llm_service.invoke(outline_prompt)

    try:
        outline_json = extract_json(outline_text)
    except Exception:
        match = re.search(r"\{.*\}", outline_text, re.DOTALL)
        if match:
            try:
                outline_json = json.loads(match.group(0))
            except Exception:
                return None
        else:
            return None

    course = Course(
        user_id=user.id,
        role=role,
        title=outline_json.get("course_title", f"{level.title()} {role} - {action.title()} Course"),
        description=outline_json.get("description", ""),
        level=level,
        status="generated",
    )
    db.add(course)
    db.flush()

    outline_modules = outline_json.get("modules", [])
    if not isinstance(outline_modules, list) or len(outline_modules) == 0:
        outline_modules = [{"module_title": t, "topics": [t]} for t in course_topics[:3]]

    total_modules = len(outline_modules)
    for idx, mod_data in enumerate(outline_modules):
        title = mod_data.get("module_title") or mod_data.get("title") or f"Module {idx+1}"
        description = (
            ", ".join(mod_data.get("topics", []))
            if mod_data.get("topics")
            else mod_data.get("description", "")
        )
        module = Module(
            course_id=course.id,
            title=title,
            description=description,
            order_index=idx,
            is_unlocked=(idx == 0),
            is_final=(idx == total_modules - 1),
        )
        db.add(module)
        db.flush()

    db.commit()
    return course.id


@app.get("/course/{course_id}", response_class=HTMLResponse)
def course_page(request: Request, course_id: int, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")

    course = db.query(Course).filter(Course.id == course_id).first()
    if not course:
        raise HTTPException(status_code=404, detail="Course not found")
    if course.user_id != user.id:
        raise HTTPException(status_code=403, detail="Not authorized")

    modules = (
        db.query(Module)
        .filter(Module.course_id == course.id)
        .order_by(Module.order_index.asc())
        .all()
    )

    return templates.TemplateResponse(request, "course.html", {
            "request": request,
            "username": user.username,
            "course": course,
            "modules": modules,
        },
    )


@app.get("/module/{module_id}", response_class=HTMLResponse)
def module_page(request: Request, module_id: int, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")

    module = db.query(Module).filter(Module.id == module_id).first()
    if not module:
        raise HTTPException(status_code=404, detail="Module not found")

    course = db.query(Course).filter(Course.id == module.course_id).first()
    if not course or course.user_id != user.id:
        raise HTTPException(status_code=403, detail="Not authorized")

    if not module.is_unlocked:
        raise HTTPException(status_code=403, detail="Module not unlocked yet")

    return templates.TemplateResponse(request, "module.html", {
            "request": request,
            "username": user.username,
            "module_id": module.id,
            "module_title": module.title,
            "course_id": course.id,
            "role": course.role,
            "level": course.level,
            "is_final": module.is_final,
        },
    )

@app.get("/api/module/{module_id}")
async def get_module(module_id: int, request: Request, db: Session = Depends(get_db)):

    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")

    module = db.query(Module).filter(Module.id == module_id).first()

    if not module:
        raise HTTPException(status_code=404, detail="Module not found")

    def _extract_practice_links(content_text: str):
        marker_start = "<!-- MODULE_PRACTICE_LINKS_START"
        marker_end = "MODULE_PRACTICE_LINKS_END -->"
        if marker_start in content_text and marker_end in content_text:
            try:
                payload = content_text.split(marker_start, 1)[1].split(marker_end, 1)[0].strip()
                parsed = json.loads(payload)
                links = parsed.get("practice_links", []) if isinstance(parsed, dict) else []
                clean_content = content_text.split(marker_start, 1)[0].strip()
                return clean_content, links
            except Exception:
                return content_text, []
        return content_text, []

    # Check if module is unlocked
    if not module.is_unlocked:
        raise HTTPException(status_code=403, detail="Module locked")

    # ✅ RETURN CACHED CONTENT
    if module.content:
        logger.info("Using cached module")
        content_text, cached_links = _extract_practice_links(module.content)
        return {
            "module_id": module.id,
            "module_title": module.title,
            "content_markdown": content_text,
            "quiz": module.quiz if isinstance(module.quiz, list) else [],
            "practice_links": cached_links,
        }

    course = db.query(Course).filter(Course.id == module.course_id).first()


    prompt = prompt_manager.get_prompt(
        "course_module_detail",
        skill=course.role,
        module=module.title,
        level=course.level,
        is_final=module.is_final,
        previous_modules=""
    )

    raw = await llm_service.invoke(prompt)

    try:
        module_json = extract_json(raw)
    except Exception as e:
        logger.error(f"Module parse error for {module_id}: {str(e)}")
        logger.error(f"Raw LLM response: {repr(raw)}")
        raise HTTPException(
            status_code=500,
            detail="Failed to parse module content from LLM"
        )

    raw_content = module_json.get("content_markdown", "") or "Detailed module content could not be retrieved. Please refresh the module."
    raw_content, _ = _extract_practice_links(raw_content)
    module.content = raw_content
    quiz_data = module_json.get("quiz", [])
    if isinstance(quiz_data, list) and len(quiz_data) == 3:
        valid = True
        for q in quiz_data:
            if not isinstance(q, dict):
                valid = False
                break
            options = q.get("options", [])
            answer = (q.get("answer") or "").upper()
            if not isinstance(options, list) or len(options) != 4 or answer not in ["A", "B", "C", "D"]:
                valid = False
                break
        if not valid:
            quiz_data = []
    else:
        quiz_data = []
    module.quiz = quiz_data

    practice_links = module_json.get("practice_links", [])
    if not isinstance(practice_links, list):
        practice_links = []

    validated_links = []
    for item in practice_links:
        if not isinstance(item, dict):
            continue
        title = item.get("title")
        url = item.get("url")
        if isinstance(title, str) and title and isinstance(url, str) and url.startswith("http"):
            validated_links.append({"title": title, "url": url})

    if not validated_links:
        external_links = module_json.get("external_practice_tasks", [])
        if isinstance(external_links, list):
            for item in external_links:
                if not isinstance(item, dict):
                    continue
                title = item.get("title")
                url = item.get("url")
                if isinstance(title, str) and title and isinstance(url, str) and url.startswith("http"):
                    validated_links.append({"title": title, "url": url})

    if len(validated_links) > 4:
        validated_links = validated_links[:4]

    if len(validated_links) < 2:
        validated_links = []

    if validated_links:
        marker = f"\n\n<!-- MODULE_PRACTICE_LINKS_START\n{json.dumps({'practice_links': validated_links})}\nMODULE_PRACTICE_LINKS_END -->"
        module.content = module.content.strip() + marker

    db.commit()
    logger.info("Module generated")

    return {
        "module_id": module.id,
        "module_title": module.title,
        "content_markdown": raw_content,
        "quiz": module.quiz,
        "practice_links": validated_links
    }

@app.post("/api/module/{module_id}/submit")
async def submit_quiz(
    module_id: int,
    request: Request,
    payload: dict = Body(...),
    db: Session = Depends(get_db)
):
    """
    Submit quiz answers for a module.
    """
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")

    try:
        # Fetch module
        module = db.query(Module).filter(Module.id == module_id).first()
        if not module:
            raise HTTPException(status_code=404, detail="Module not found")

        # Check authorization
        course = db.query(Course).filter(Course.id == module.course_id).first()
        if course.user_id != user.id:
            raise HTTPException(status_code=403, detail="Not authorized")

        # Ensure quiz is loaded
        if not module.quiz or not isinstance(module.quiz, list) or len(module.quiz) != 3:
            raise HTTPException(status_code=400, detail="Module has no valid quiz")

        user_answers = payload.get("answers", [])
        if not isinstance(user_answers, list) or len(user_answers) != 3:
            raise HTTPException(status_code=400, detail="Must provide exactly 3 answers")

        # Score the quiz
        score = 0
        for user_ans, quiz_item in zip(user_answers, module.quiz):
            correct_ans = (quiz_item.get("answer") or "").upper()
            if isinstance(user_ans, str) and user_ans.upper() == correct_ans:
                score += 1

        passed = score >= 2  # At least 2 out of 3 correct
        # Store or update attempt
        attempt = (
            db.query(ModuleAttempt)
            .filter(
                ModuleAttempt.user_id == user.id,
                ModuleAttempt.module_id == module_id
            )
            .order_by(ModuleAttempt.created_at.desc())
            .first()
        )
        attempt_count = 1
        if attempt:
            existing_payload = attempt.answers if isinstance(attempt.answers, dict) else {}
            try:
                previous_count = int(existing_payload.get("attempt_count", 1))
            except Exception:
                previous_count = 1
            attempt_count = previous_count + 1
            attempt.score = score
            attempt.total_questions = 3
            attempt.answers = {
                "answers": user_answers,
                "is_passed": passed,
                "attempt_count": attempt_count,
            }
        else:
            attempt = ModuleAttempt(
                user_id=user.id,
                module_id=module_id,
                score=score,
                total_questions=3,
                answers={
                    "answers": user_answers,
                    "is_passed": passed,
                    "attempt_count": attempt_count,
                }
            )
            db.add(attempt)
        logger.info(
            "Module attempt persisted user_id=%s module_id=%s score=%s passed=%s attempts=%s",
            user.id,
            module_id,
            score,
            passed,
            attempt_count,
        )

        next_module_id = None
        # Mark module as completed and unlock next module if passed
        if passed:
            module.is_completed = True
            next_module = db.query(Module).filter(
                Module.course_id == module.course_id,
                Module.order_index == module.order_index + 1
            ).first()
            if next_module:
                next_module.is_unlocked = True
                next_module_id = next_module.id
                logger.info("Next module unlocked user_id=%s module_id=%s", user.id, next_module_id)

        db.commit()

        logger.info("Quiz submitted")

        return {
            "score": score,
            "passed": passed,
            "is_final": module.is_final,
            "next_module_id": next_module_id,
            "role": course.role,
            "level": course.level,
            "course_id": course.id
        }

    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"Quiz submission failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Quiz submission failed: {str(e)}")


@app.get("/api/course/{course_id}/status")
async def get_course_status(
    course_id: int,
    request: Request,
    db: Session = Depends(get_db)
):
    """
    Get course completion status.
    
    - Check if all modules completed
    - Return: modules status, interview_unlocked flag
    """
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")
    
    try:
        # Fetch course
        course = db.query(Course).filter(Course.id == course_id).first()
        if not course:
            raise HTTPException(status_code=404, detail="Course not found")
        
        # Check authorization
        if course.user_id != user.id:
            raise HTTPException(status_code=403, detail="Not authorized")
        
        # Get all modules
        modules = db.query(Module).filter(Module.course_id == course_id).order_by(Module.order_index).all()
        
        module_status = []
        all_completed = True
        for mod in modules:
            # Get best attempt score for this module
            best_attempt = db.query(ModuleAttempt).filter(
                ModuleAttempt.user_id == user.id,
                ModuleAttempt.module_id == mod.id
            ).order_by(ModuleAttempt.score.desc()).first()
            
            module_status.append({
                "id": mod.id,
                "title": mod.title,
                "order_index": mod.order_index,
                "is_unlocked": mod.is_unlocked,
                "is_completed": mod.is_completed,
                "best_score": best_attempt.score if best_attempt else None,
                "best_percentage": int((best_attempt.score / best_attempt.total_questions) * 100) if best_attempt else None
            })
            
            if not mod.is_completed:
                all_completed = False
        
        return {
            "course_id": course_id,
            "title": course.title,
            "role": course.role,
            "level": course.level,
            "modules": module_status,
            "all_completed": all_completed,
            "interview_unlocked": all_completed  # Unlock interview after all modules
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Course status check failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Course status check failed: {str(e)}")


async def generate_final_interview(
    user_id: int,
    course_id: int,
    db: Session
) -> dict:
    """
    Generate a final interview session after course completion.
    
    Uses:
    - Course role and level
    - User's resume
    - Weak topics (modules with low scores)
    
    Returns:
    - interview_session_id
    - initial questions
    """
    try:
        # Fetch course and user info
        course = db.query(Course).filter(Course.id == course_id).first()
        if not course:
            raise ValueError("Course not found")
        
        user_profile = db.query(UserProfile).filter(UserProfile.user_id == user_id).first()
        
        # Identify weak topics (modules with score < 70%)
        modules = db.query(Module).filter(Module.course_id == course_id).all()
        weak_topics = []
        
        for mod in modules:
            best_attempt = db.query(ModuleAttempt).filter(
                ModuleAttempt.user_id == user_id,
                ModuleAttempt.module_id == mod.id
            ).order_by(ModuleAttempt.score.desc()).first()
            
            if best_attempt:
                percentage = (best_attempt.score / best_attempt.total_questions) * 100
                if percentage < 70:
                    weak_topics.append(mod.title)
        
        # Create interview session
        session = InterviewSession(
            user_id=user_id,
            role=course.role,
            level=course.level,
            status="active"
        )
        db.add(session)
        db.flush()
        
        # Generate initial questions using interview chain
        resume_text = user_profile.resume_file_path if user_profile else ""
        
        initial_questions = await question_chain.invoke({
            "role": course.role,
            "level": course.level,
            "count": 5,
            "resume_text": resume_text,
            "previous_questions": [],
            "used_categories": []
        })
        
        db.commit()
        
        return {
            "interview_session_id": session.id,
            "role": course.role,
            "level": course.level,
            "weak_topics": weak_topics,
            "status": "started"
        }
    
    except Exception as e:
        db.rollback()
        logger.error(f"Final interview generation failed: {str(e)}")
        raise


@app.post("/api/course/{course_id}/start-interview")
async def start_final_interview(
    course_id: int,
    request: Request,
    db: Session = Depends(get_db)
):
    """
    Trigger final interview generation after course completion.
    
    - Check all modules completed
    - Generate interview session
    - Return session ID
    """
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")
    
    try:
        # Verify course exists and belongs to user
        course = db.query(Course).filter(
            Course.id == course_id,
            Course.user_id == user.id
        ).first()
        if not course:
            raise HTTPException(status_code=404, detail="Course not found")
        
        # Check if all modules completed
        modules = db.query(Module).filter(Module.course_id == course_id).all()
        all_completed = all(mod.is_completed for mod in modules)
        
        if not all_completed:
            raise HTTPException(status_code=400, detail="Not all modules completed")
        
        # Generate final interview
        interview_data = await generate_final_interview(user.id, course_id, db)
        
        return interview_data
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Interview start failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Interview start failed: {str(e)}")


# ===========================================================================
# RAG MONITORING ENDPOINTS
# ===========================================================================

@app.get("/api/rag/stats")
async def get_rag_stats():
    """Get RAG pipeline statistics and performance metrics."""
    
    if rag_pipeline is None:
        return JSONResponse(
            {"error": "RAG pipeline not initialized"},
            status_code=503
        )
    
    try:
        stats = rag_pipeline.get_stats()
        return {
            "status": "ok",
            "initialized": stats["initialized"],
            "documents_indexed": stats["vector_store"]["total_documents"],
            "retrievals_performed": stats["retrieval_stats"]["retrievals_performed"],
            "cache_hit_rate": (
                stats["retrieval_stats"]["cache_hits"] /
                (stats["retrieval_stats"]["cache_hits"] + stats["retrieval_stats"]["cache_misses"])
                if (stats["retrieval_stats"]["cache_hits"] + stats["retrieval_stats"]["cache_misses"]) > 0
                else 0
            ),
            "avg_retrieval_time_ms": stats["retrieval_stats"]["avg_retrieval_time"] * 1000,
            "cache_size": stats["cache_size"],
            "embedding_dimension": stats["vector_store"]["embedding_dimension"],
            "document_categories": stats["ingestion_stats"]["categories"]
        }
    
    except Exception as e:
        return JSONResponse(
            {"error": str(e)},
            status_code=500
        )


@app.post("/api/rag/clear-cache")
async def clear_rag_cache():
    """Clear the RAG retrieval cache."""
    
    if rag_pipeline is None:
        return JSONResponse(
            {"error": "RAG pipeline not initialized"},
            status_code=503
        )
    
    try:
        rag_pipeline.clear_cache()
        return {"status": "ok", "message": "RAG cache cleared"}
    
    except Exception as e:
        return JSONResponse(
            {"error": str(e)},
            status_code=500
        )


@app.get("/api/rag/test-retrieval")
async def test_rag_retrieval(query: str = "Python junior level interview"):
    """Test RAG retrieval with a sample query (for debugging)."""
    
    if rag_pipeline is None or not rag_pipeline.initialized:
        return JSONResponse(
            {"error": "RAG pipeline not initialized"},
            status_code=503
        )
    
    try:
        from services.rag.rag_config import RetrievalContext
        
        context = await rag_pipeline.retrieve_context(query)
        
        return {
            "status": "ok",
            "query": query,
            "retrieved_context": context[:500] + "..." if len(context) > 500 else context
        }
    
    except Exception as e:
        return JSONResponse(
            {"error": str(e)},
            status_code=500
        )